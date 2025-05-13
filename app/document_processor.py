"""Document processor module for handling different document types."""

import aiofiles
from typing import List, Union, Dict, Any
from pathlib import Path
import json
import pandas as pd
from pypdf import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

class DocumentProcessor:
    """Handles the processing of different document types."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )
    
    async def process_file(self, file_path: Union[str, Path], file_type: str) -> List[LangchainDocument]:
        """Process a file based on its type asynchronously.
        
        Args:
            file_path: Path to the file
            file_type: Type of the file (pdf, docx, txt, json, xlsx, html)
            
        Returns:
            List of Langchain Document objects
        """
        file_path = Path(file_path)
        content = ""
        
        if file_type == "pdf":
            content = await self._process_pdf(file_path)
        elif file_type == "docx":
            content = await self._process_docx(file_path)
        elif file_type == "txt":
            content = await self._process_text(file_path)
        elif file_type == "json":
            content = await self._process_json(file_path)
        elif file_type == "xlsx":
            content = await self._process_excel(file_path)
        elif file_type == "html":
            content = await self._process_html(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return await self._create_documents(content, file_path)
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files asynchronously, including OCR if needed."""
        # PDF reading is not async, but we'll wrap it for consistency
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if not page_text.strip():  # If page is empty, try OCR
                if hasattr(page, "images"):
                    for image in page.images:
                        img = Image.open(image)
                        text += pytesseract.image_to_string(img) + "\n"
            else:
                text += page_text + "\n"
        return text
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process Word documents asynchronously."""
        # Document reading is not async, but we'll wrap it for consistency
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    async def _process_text(self, file_path: Path) -> str:
        """Process text files asynchronously."""
        async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
            return await f.read()
    
    async def _process_json(self, file_path: Path) -> str:
        """Process JSON files asynchronously."""
        async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
            content = await f.read()
            data = json.loads(content)
            return json.dumps(data, indent=2)
    
    async def _process_excel(self, file_path: Path) -> str:
        """Process Excel files asynchronously."""
        # Pandas read_excel is not async, but we'll wrap it for consistency
        df = pd.read_excel(file_path)
        return df.to_string()
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML files asynchronously."""
        async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
            content = await f.read()
            soup = BeautifulSoup(content, "html.parser")
            return soup.get_text()
    
    async def _create_documents(self, content: str, source: Path) -> List[LangchainDocument]:
        """Create Langchain documents from content asynchronously."""
        # Text splitting is not async, but we'll wrap it for consistency
        texts = self.text_splitter.split_text(content)
        return [
            LangchainDocument(
                page_content=text,
                metadata={"source": str(source)}
            )
            for text in texts
        ] 